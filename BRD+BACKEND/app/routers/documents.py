from fastapi import APIRouter, File, UploadFile, Form, Depends, HTTPException
from typing import List
from sqlalchemy.orm import Session
from app import models, schemas, auth
from app.v2_logic import process_single_document
from app.services.file_storage import LocalFileStorage
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import tempfile

router = APIRouter()
storage = LocalFileStorage()

def extract_text(file_bytes: bytes, content_type: str) -> str:
    try:
        if content_type == "application/pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file_bytes)
                tmp.flush()
                reader = PdfReader(tmp.name)
                return "\n".join(page.extract_text() or "" for page in reader.pages)

        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(file_bytes)
                tmp.flush()
                doc = DocxDocument(tmp.name)
                return "\n".join([para.text for para in doc.paragraphs])

        elif content_type.startswith("text/"):
            return file_bytes.decode("utf-8", errors="ignore")

        return "Unsupported file format"
    except Exception as e:
        return f"Failed to extract file content: {e}"

@router.post("/upload", response_model=schemas.DocumentOut)
async def upload_support_file(
    file: UploadFile = File(...),
    description: str = Form(...),
    project_id: int = Form(...),
    db: Session = Depends(auth.get_db),
    current_user: models.User = Depends(auth.get_current_user)
):
    project = db.query(models.Project).filter(models.Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Invalid project ID")

    file_bytes = await file.read()
    saved_path = storage.save(file_bytes, file.filename)

    extracted_text = extract_text(file_bytes, file.content_type)
    print(f"Extracted text: {extracted_text}")
    result = process_single_document(
        file_content=extracted_text.encode("utf-8"),
        description=description,
        document_type=file.content_type,
        session_folder="processed_docs"
    )

    summary = result["summary"]
    doc_db = models.Document(
        filename=file.filename,
        path=saved_path,
        description=description,
        summary_title=summary.get("title"),
        summary_description=summary.get("description"),
        project_id=project_id,
        uploaded_by=current_user.id
    )
    db.add(doc_db)
    db.commit()
    db.refresh(doc_db)

    return doc_db

@router.get("/project/{project_id}", response_model=List[schemas.DocumentOut])
def list_documents(project_id: int, db: Session = Depends(auth.get_db), current_user: models.User = Depends(auth.get_current_user)):
    docs = db.query(models.Document).filter(models.Document.project_id == project_id).all()
    return docs
