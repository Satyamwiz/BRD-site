import os
import json
import re
import uuid
import datetime
import logging
from typing import List, Dict, Any, Union
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import requests
from groq import Groq

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

class SupportDocument:
    def __init__(self, file_content: bytes, description: str, document_type: str):
        self.file_content = file_content
        self.description = description
        self.document_type = document_type

class BRDInput:
    def __init__(self, prompt: str, template: bytes, support_documents: List[SupportDocument]):
        self.prompt = prompt
        self.template = template
        self.support_documents = support_documents

class RewordSummaryAgent:
    def process(self, brd_input: BRDInput) -> str:
        prompt = f"""
        You are an expert Business Analyst. 
        Your task is to carefully analyze the provided content, 
        which may include email chains, meeting notes, or transcriptions. 
        Based on this analysis, create a comprehensive, concise and structured writeup. 
        This writeup will serve as input notes for the detailed preparation of a Business Requirements Document (BRD). 
        Ensure that your writeup captures all key points, decisions, and action items relevant to the project or business process.

        {brd_input.prompt}
        Your task is to generate a JSON object with the following structure:
        {{
          "title": "One line title summarizing the main topic",
          "description": "A brif writeup"
        }}

        Guidelines for the response:
        1. The response must be a valid JSON object.
        2. Do not include any text before or after the JSON object.
        3. The "title" should be a concise, one-line summary.
        4. The "description" should be extensive, typically several paragraphs long, covering all aspects of the analyzed content in detail.
        5. Include all key points, decisions, action items, and relevant information from the document in the description.
        6. Ensure proper JSON formatting, including using double quotes for strings and escaping any special characters.

        Remember, your entire response should be a single, valid JSON object.

        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content
    
class BRDCompletionAgent:
    def ensure_string(self, data: Union[str, bytes]) -> str:
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                for encoding in ['iso-8859-1', 'windows-1252']:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return data.decode('latin-1', errors='replace')
        return data

    def process(self, brd_input: BRDInput, reworded_summary: str) -> Dict[str, Any]:
        template = self.ensure_string(brd_input.template)
        prompt=f"""As an experienced project manager, analyze the following BRD summary and determine if additional details are needed:

        {reworded_summary}

        Use the following template to structure your BRD:

        {template}

        Consider the following aspects:
        1. Project scope: Is it clearly defined?
        2. Stakeholder requirements: Are all key stakeholders' needs addressed?
        3. Technical specifications: Are there enough details for the development team?
        4. Timeline and milestones: Are they mentioned or need to be clarified?
        5. Budget considerations: Is there any mention of budget constraints or requirements?
        6. Risk assessment: Are potential risks identified or do they need to be addressed?
        7. Success criteria: Are the project's success metrics clearly defined?

        Based on your analysis, respond with a JSON object in the following format:
        {{
            "status": "need" or "not_need",
            "details": [] (empty list if status is "not_need", otherwise list of required details)
        }}

        If "need" status is determined, provide specific questions or points that need clarification in the "details" list.

        Instructions:
        1. Ensure your entire response is a valid JSON object.
        2. Do not include any text before or after the JSON object.
        3. Use double quotes for all strings in the JSON.
        4. The "status" field must be exactly "need" or "not_need".
        5. The "details" field must be a list, even if it's empty.
        6. Do not include any explanations or additional text outside the JSON structure.
        7. Verify that your response can be parsed as JSON before submitting.
        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        try:
            return json.loads(response.choices[0].message.content)
        except json.JSONDecodeError:
            return {"status": "error", "details": ["Failed to parse response"]}

class BRDCreationAgent:
    def ensure_string(self, data: Union[str, bytes]) -> str:
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                for encoding in ['iso-8859-1', 'windows-1252']:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return data.decode('latin-1', errors='replace')
        return data

    def process(self, brd_input: BRDInput, reworded_summary: str, completion_suggestions: Dict[str, Any]) -> str:
        template = self.ensure_string(brd_input.template)
        prompt = f"""
        As a senior business analyst, create a comprehensive Business Requirements Document (BRD) based on the following information:

        Initial Summary:
        {reworded_summary}

        Additional Details:
        {json.dumps(completion_suggestions, indent=2)}

        Use the following template to structure your BRD:

        {template}

        Ensure that you:
        - Refer to these instructions and ask for any missing information related to them.
        - Address all points mentioned in the initial summary and additional details.
        - Provide clear, actionable requirements using industry-standard terminology and formatting.
        - Fill in the document with appropriate content for each section, using placeholders for any information that may need to be filled in later.
        - Avoid adding any extra text, such as introductory statements.
        - Address all points mentioned in the initial summary and additional details
        - Provide clear, actionable requirements
        - Include placeholders for any information that might need to be filled in later
        - do not add any extra text like (Here is a comprehensive Business Requirements Document (BRD) for the Life Secure Flexi Plan Insurance Product Development and CRM System Integration project:)

        Create a detailed and professional BRD that could be presented to senior management and technical teams.
        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content

class BRDReviewAgent:
    def process(self, brd_document: str) -> str:
        prompt = f"""
        As a senior project manager and business analyst, review the following Business Requirements Document (BRD) and provide comprehensive feedback:

        {brd_document}

        In your review, please consider the following aspects:
        1. Completeness: Are all necessary sections present and adequately detailed?
        2. Clarity: Is the document easy to understand for both technical and non-technical stakeholders?
        3. Consistency: Are there any contradictions or inconsistencies in the requirements?
        4. Feasibility: Do the requirements seem realistic and achievable?
        5. Traceability: Can each requirement be traced back to a business need or stakeholder request?
        6. Testability: Are the requirements specific enough to be tested?
        7. Prioritization: Is there a clear indication of which requirements are must-haves vs. nice-to-haves?
        8. Risks: Are potential risks adequately identified and addressed?
        9. Compliance: Does the document adhere to any relevant industry standards or regulations?
        10. Overall quality: Does the document meet the standards expected of a professional BRD?

        Provide detailed feedback, including:
        - Strengths of the document
        - Areas for improvement
        - Specific suggestions for enhancing the BRD
        - Any critical omissions or concerns

        Your review should be thorough and constructive, aimed at improving the overall quality and effectiveness of the BRD.
        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content

# ------------ Logic Wrappers -------------

def process_brd(brd_input: BRDInput) -> Dict[str, Any]:
    reword_agent = RewordSummaryAgent()
    completion_agent = BRDCompletionAgent()
    creation_agent = BRDCreationAgent()

    reworded_summary = reword_agent.process(brd_input)
    completion_suggestions = completion_agent.process(brd_input, reworded_summary)
    brd_draft = creation_agent.process(brd_input, reworded_summary, completion_suggestions)

    return {
        "reworded_summary": reworded_summary,
        "completion_suggestions": completion_suggestions,
        "brd_draft": brd_draft
    }

def generate_final_brd(brd_input: BRDInput, completion_answers: Dict[str, str], reworded_summary: str) -> Dict[str, Any]:
    creation_agent = BRDCreationAgent()
    review_agent = BRDReviewAgent()

    completion_suggestions = {
        "status": "need",
        "details": list(completion_answers.keys())
    }

    brd_document = creation_agent.process(brd_input, reworded_summary, completion_suggestions)
    review_feedback = review_agent.process(brd_document)

    return {
        "brd_document": brd_document,
        "review_feedback": review_feedback
    }

# ------------ DOCX Export -------------

def create_brd_word_document(brd_content: str, output_path: str):
    doc = Document()
    styles = doc.styles
    if 'BRDNormal' not in styles:
        style = styles.add_style('BRDNormal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(11)
    doc.add_paragraph(brd_content, style='BRDNormal')
    doc.save(output_path)


def process_single_document(file_content: bytes, description: str, document_type: str, session_folder: str) -> Dict[str, Any]:
    reword_agent = RewordSummaryAgent()
    support_doc = SupportDocument(file_content, description, document_type)
    brd_input = BRDInput(prompt=f"Summarize the following document: {description}", template=b"", support_documents=[support_doc])

    try:
        reworded_summary = reword_agent.process(brd_input)
        logger.info("Reworded summary processed")

        try:
            summary_json = json.loads(reworded_summary)
        except json.JSONDecodeError:
            logger.warning("Response is not in JSON format. Creating a default JSON structure.")
            summary_json = {
                "title": description[:50],
                "description": reworded_summary
            }

        os.makedirs(session_folder, exist_ok=True)
        unique_id = str(uuid.uuid4())[:8]
        filename = re.sub(r'[^\w\-_\. ]', '', summary_json.get("title", "untitled")).replace(' ', '_')
        file_path = os.path.join(session_folder, f"{filename}_{unique_id}.json")
        with open(file_path, "w") as f:
            json.dump(summary_json, f, indent=2)

        return {"summary": summary_json, "path": file_path}

    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return {"summary": {"title": "Error", "description": str(e)}, "path": ""}
