import groq
from typing import List, Dict, Any, Union
import io
import json
import streamlit as st
import os
import re
import markdown
import logging
import uuid
import datetime
import docx
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import requests
from io import BytesIO


def add_html_to_docx(html_content, doc):
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.descendants:
        # Handling Headings (h1 to h6)
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            doc.add_heading(element.get_text(), level=level)
        
        # Handling Paragraphs
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        
        # Handling Bold and Italic
        elif element.name == 'strong':
            bold_para = doc.add_paragraph()
            bold_para.add_run(element.get_text()).bold = True
        elif element.name == 'em':
            italic_para = doc.add_paragraph()
            italic_para.add_run(element.get_text()).italic = True
        
        # Handling Unordered Lists
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(f'- {li.get_text()}', style='ListBullet')
        
        # Handling Ordered Lists
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(f'{li.get_text()}', style='ListNumber')
        
        # Handling Links
        elif element.name == 'a' and element.get('href'):
            link_text = element.get_text()
            link_url = element['href']
            link_para = doc.add_paragraph()
            link_run = link_para.add_run(link_text)
            link_run.font.underline = True  # Adding underline to indicate it's a link
            link_run.font.color.rgb = (0, 0, 255)  # Blue color for links
            # (Word doesn't support clickable links by default in `python-docx`)
        
        # Handling Images
        elif element.name == 'img' and element.get('src'):
            try:
                image_url = element['src']
                image_response = requests.get(image_url)
                image_data = BytesIO(image_response.content)
                doc.add_picture(image_data, width=Inches(2))  # Resizing to 2 inches
            except Exception as e:
                print(f"Failed to add image: {e}")
        
        # Handling Tables
        elif element.name == 'table':
            rows = element.find_all('tr')
            table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all('td')))
            for row_idx, row in enumerate(rows):
                cells = row.find_all('td')
                for col_idx, cell in enumerate(cells):
                    table.cell(row_idx, col_idx).text = cell.get_text()

def html_to_word(html_content, output_file):
    # Ensure that doc is initialized correctly
    try:
        doc = Document()  # Initialize the Document
    except Exception as e:
        print(f"Error initializing Document: {e}")
        return  # If failed, exit the function

    # Add HTML content to the document
    try:
        doc = add_html_to_docx(html_content, doc)  # Make sure this returns a valid doc object
    except Exception as e:
        print(f"Error while adding HTML content: {e}")
        return  # If failed, exit the function

    # Save the document
    try:
        doc.save(output_file)  # Save to output file
    except Exception as e:
        print(f"Error saving Document: {e}")


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Groq client setup
groq_client = groq.Groq(api_key="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")

class SupportDocument:
    def __init__(self, file_content: bytes, description: str, document_type: str):
        self.file_content = file_content
        self.description = description
        self.document_type = document_type

class BRDInput:
    def __init__(self, prompt: str, template: bytes, support_documents: List[SupportDocument]):
        self.prompt = prompt
        self.template = template
        self.support_documents = support_documents

class RewordSummaryAgent:
    def process(self, brd_input: BRDInput) -> str:
        prompt = f"""
        You are an expert Business Analyst. 
        Your task is to carefully analyze the provided content, 
        which may include email chains, meeting notes, or transcriptions. 
        Based on this analysis, create a comprehensive, concise and structured writeup. 
        This writeup will serve as input notes for the detailed preparation of a Business Requirements Document (BRD). 
        Ensure that your writeup captures all key points, decisions, and action items relevant to the project or business process.

        {brd_input.prompt}
        Your task is to generate a JSON object with the following structure:
        {{
          "title": "One line title summarizing the main topic",
          "description": "A brif writeup"
        }}

        Guidelines for the response:
        1. The response must be a valid JSON object.
        2. Do not include any text before or after the JSON object.
        3. The "title" should be a concise, one-line summary.
        4. The "description" should be extensive, typically several paragraphs long, covering all aspects of the analyzed content in detail.
        5. Include all key points, decisions, action items, and relevant information from the document in the description.
        6. Ensure proper JSON formatting, including using double quotes for strings and escaping any special characters.

        Remember, your entire response should be a single, valid JSON object.

        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content

class BRDCompletionAgent:

    def ensure_string(self, data: Union[str, bytes]) -> str:
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                # If UTF-8 fails, try other common encodings
                for encoding in ['iso-8859-1', 'windows-1252']:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                # If all else fails, use latin-1 as a last resort
                return data.decode('latin-1', errors='replace')
        elif isinstance(data, str):
            return data
        else:
            raise ValueError(f"Unsupported data type for template: {type(data)}")
        
    def process(self, brd_input: BRDInput, reworded_summary: str) -> Dict[str, Any]:
        template = self.ensure_string(brd_input.template)
        prompt = f"""
        As an experienced project manager, analyze the following BRD summary and determine if additional details are needed:

        {reworded_summary}

        Use the following template to structure your BRD:

        {template}

        Consider the following aspects:
        1. Project scope: Is it clearly defined?
        2. Stakeholder requirements: Are all key stakeholders' needs addressed?
        3. Technical specifications: Are there enough details for the development team?
        4. Timeline and milestones: Are they mentioned or need to be clarified?
        5. Budget considerations: Is there any mention of budget constraints or requirements?
        6. Risk assessment: Are potential risks identified or do they need to be addressed?
        7. Success criteria: Are the project's success metrics clearly defined?

        Based on your analysis, respond with a JSON object in the following format:
        {{
            "status": "need" or "not_need",
            "details": [] (empty list if status is "not_need", otherwise list of required details)
        }}

        If "need" status is determined, provide specific questions or points that need clarification in the "details" list.

        Instructions:
        1. Ensure your entire response is a valid JSON object.
        2. Do not include any text before or after the JSON object.
        3. Use double quotes for all strings in the JSON.
        4. The "status" field must be exactly "need" or "not_need".
        5. The "details" field must be a list, even if it's empty.
        6. Do not include any explanations or additional text outside the JSON structure.
        7. Verify that your response can be parsed as JSON before submitting.
        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        
        # Print the raw response content
        print("Raw API response:")
        print(response.choices[0].message.content)
        
        try:
            return json.loads(response.choices[0].message.content)
        except json.JSONDecodeError as e:
            print(f"Error decoding JSON: {e}")
            return {"status": "error", "details": ["Failed to parse API response"]}
        

class BRDCreationAgent:
    def ensure_string(self, data: Union[str, bytes]) -> str:
        if isinstance(data, bytes):
            try:
                return data.decode('utf-8')
            except UnicodeDecodeError:
                # If UTF-8 fails, try other common encodings
                for encoding in ['iso-8859-1', 'windows-1252']:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                # If all else fails, use latin-1 as a last resort
                return data.decode('latin-1', errors='replace')
        elif isinstance(data, str):
            return data
        else:
            raise ValueError(f"Unsupported data type for template: {type(data)}")
    def process(self, brd_input: BRDInput, reworded_summary: str, completion_suggestions: Dict[str, Any]) -> str:
        try:
            template = self.ensure_string(brd_input.template)
        except ValueError as e:
            print(f"Error processing template: {str(e)}")
            template = "Unable to process template. Proceeding without it."
        prompt = f"""
        As a senior business analyst, create a comprehensive Business Requirements Document (BRD) based on the following information:

        Initial Summary:
        {reworded_summary}

        Additional Details:
        {json.dumps(completion_suggestions, indent=2)}

        Use the following template to structure your BRD:

        {template}

        Ensure that you:
        - Refer to these instructions and ask for any missing information related to them.
        - Address all points mentioned in the initial summary and additional details.
        - Provide clear, actionable requirements using industry-standard terminology and formatting.
        - Fill in the document with appropriate content for each section, using placeholders for any information that may need to be filled in later.
        - Avoid adding any extra text, such as introductory statements.
        - Address all points mentioned in the initial summary and additional details
        - Provide clear, actionable requirements
        - Include placeholders for any information that might need to be filled in later
        - do not add any extra text like (Here is a comprehensive Business Requirements Document (BRD) for the Life Secure Flexi Plan Insurance Product Development and CRM System Integration project:)

        Create a detailed and professional BRD that could be presented to senior management and technical teams.
        """
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content

class BRDReviewAgent:
    def process(self, brd_document: str) -> str:
        
        response = groq_client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model="llama3-70b-8192",
        )
        return response.choices[0].message.content

def read_file(file_path):
    with open(file_path, 'rb') as file:
        return file.read()

def process_brd(brd_input: BRDInput) -> Dict[str, Any]:
    reword_agent = RewordSummaryAgent()
    completion_agent = BRDCompletionAgent()

    reworded_summary = reword_agent.process(brd_input)
    completion_suggestions = completion_agent.process(brd_input, reworded_summary)

    return {
        "reworded_summary": reworded_summary,
        "completion_suggestions": completion_suggestions,
    }

def generate_final_brd(brd_input: BRDInput, completion_answers: Dict[str, str], reworded_summary: str) -> Dict[str, Any]:
    creation_agent = BRDCreationAgent()
    review_agent = BRDReviewAgent()

    # Prepare completion suggestions
    completion_suggestions = {
        "status": "need",
        "details": list(completion_answers.keys())
    }

    brd_document = creation_agent.process(brd_input, reworded_summary, completion_suggestions)
    review_feedback = review_agent.process(brd_document)

    return {
        "brd_document": brd_document,
        "review_feedback": review_feedback
    }

def sanitize_filename(filename):
    # Remove invalid characters and replace spaces with underscores
    return re.sub(r'[^\w\-_\. ]', '', filename).replace(' ', '_')

def ensure_directory_exists(directory):
    """Ensure the specified directory exists and is writable."""
    if not os.path.exists(directory):
        try:
            os.makedirs(directory)
            logger.info(f"Created directory: {directory}")
        except OSError as e:
            logger.error(f"Failed to create directory {directory}: {e}")
            return False
    
    if not os.access(directory, os.W_OK):
        logger.error(f"Directory {directory} is not writable")
        return False
    
    return True

def create_session_folder():
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    session_folder = os.path.join("processed_docs", f"session_{timestamp}")
    os.makedirs(session_folder, exist_ok=True)
    return session_folder



def process_single_document(file_content: bytes, description: str, document_type: str, session_folder: str) -> Dict[str, Any]:
    reword_agent = RewordSummaryAgent()
    support_doc = SupportDocument(file_content, description, document_type)
    brd_input = BRDInput(prompt=f"Summarize the following document: {description}", template=b"", support_documents=[support_doc])
    
    try:
        reworded_summary = reword_agent.process(brd_input)
        logger.info(f"Reworded summary processed")
        
        try:
            summary_json = json.loads(reworded_summary)
        except json.JSONDecodeError:
            logger.warning("Response is not in JSON format. Creating a default JSON structure.")
            summary_json = {
                "title": description[:50],  # Use first 50 characters of description as title
                "description": reworded_summary
            }
        
        # Save the processed document
        file_path = save_processed_document(summary_json, session_folder)
        
        return {"summary": summary_json, "path": file_path}
        
    except Exception as e:
        logger.error(f"Unexpected error in process_single_document: {e}")
        return {"summary": {"title": "Error", "description": str(e)}, "path": ""}

def save_processed_document(summary: Dict[str, str], session_folder: str) -> str:
    unique_id = str(uuid.uuid4())[:8]
    filename = sanitize_filename(f"{summary.get('title', 'untitled')}_{unique_id}")
    file_path = os.path.join(session_folder, f"{filename}.json")
    
    try:
        with open(file_path, "w") as f:
            json.dump(summary, f, indent=2)
        logger.info(f"Saved processed document to {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"Failed to save processed document: {e}")
        return ""


def read_processed_documents() -> List[Dict[str, str]]:
    processed_docs = []
    if not os.path.exists("processed_docs"):
        logger.warning("processed_docs directory does not exist")
        return processed_docs
    
    for filename in os.listdir("processed_docs"):
        if filename.endswith(".json"):
            file_path = os.path.join("processed_docs", filename)
            try:
                with open(file_path, "r") as f:
                    processed_docs.append(json.load(f))
                logger.info(f"Successfully read {file_path}")
            except json.JSONDecodeError as e:
                logger.error(f"Error reading {file_path}: {e}")
    return processed_docs

def display_processed_documents(processed_docs):
    st.header("Processed Documents")
    
    if not processed_docs:
        st.write("No processed documents found.")
        return

    # Create tabs for each document
    tabs = st.tabs([f"Document {i+1}" for i in range(len(processed_docs))])

    # Display each document in its own tab
    for i, (tab, doc) in enumerate(zip(tabs, processed_docs)):
        with tab:
            st.subheader(doc['summary'].get('title', f'Document {i+1}'))
            st.write(doc['summary'].get('description', 'No description available.'))

def create_brd_word_document(final_brd, output_filename):
    doc = docx.Document()

    # Set up styles
    styles = doc.styles
    normal_style = styles.add_style('BRD Normal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)

    # Add BRD Document
    doc.add_paragraph(final_brd["brd_document"], style='BRD Normal')

    # Save the document
    doc.save(output_filename)


def verify_directory_contents():
    st.subheader("Directory Contents")
    if not os.path.exists("processed_docs"):
        st.write("processed_docs directory does not exist")
    else:
        files = os.listdir("processed_docs")
        if not files:
            st.write("processed_docs directory is empty")
        else:
            st.write("Files in processed_docs:")
            for file in files:
                st.write(file)

def main():
    st.markdown(f"""
        <style>
            .header {{
                background-color: #4A4A4A;
                border-radius: 25px;
                padding: 20px;
                display: flex;
                align-items: center;
                color: white;
                font-size: 24px;
                width: 100%;
                white-space: nowrap;
            }}
            .header h2 {{
                margin: 0;
                margin-left: 20px;
            }}
        </style>
        <div class="header">
            <h3>Business Requirements Document (BRD) Creator</h3>
        </div>
    """, unsafe_allow_html=True)

    # Initialize session state (keeping your existing initializations)
    if 'session_folder' not in st.session_state:
        st.session_state.session_folder = create_session_folder()
    if 'processed_docs' not in st.session_state:
        st.session_state.processed_docs = []
    if 'template_file' not in st.session_state:
        st.session_state.template_file = None
    if 'support_docs' not in st.session_state:
        st.session_state.support_docs = []
    if 'processed_doc_paths' not in st.session_state:
        st.session_state.processed_doc_paths = []
    if 'processing_stage' not in st.session_state:
        st.session_state.processing_stage = 'input'
    if 'completion_answers' not in st.session_state:
        st.session_state.completion_answers = {}
    if 'brd_results' not in st.session_state:
        st.session_state.brd_results = None
    if 'final_brd' not in st.session_state:
        st.session_state.final_brd = None
    if 'need_completion' not in st.session_state:
        st.session_state.need_completion = True

    # Input section
    if st.session_state.processing_stage == 'input':
        # Custom CSS for smaller header
        # st.markdown(f"""
        #     <style>
        #         .small-header {{
        #             font-size: 20px;
        #             font-weight: bold;
        #             color: white;
        #         }}
        #     </style>
        #     <div class="small-header">
        #         Input Information
        #     </div>
        # """, unsafe_allow_html=True)
        st.markdown(f""" """, unsafe_allow_html=True)

        prompt = st.text_area("Requirment Description", "Describe what you want to include in your BRD...")
        
        # Template file upload
        template_file = st.file_uploader("Upload BRD template (PDF)", type="pdf")
        if template_file is not None:
            st.session_state.template_file = template_file
            st.success(f"Template uploaded: {template_file.name}")

        # Multiple document upload section
        st.markdown(f"""
            <style>
                .small-header {{
                    font-size: 20px;
                    font-weight: bold;
                    color: white;
                }}
            </style>
            <div class="small-header">
                Upload Requirment Specification Files
            </div>
        """, unsafe_allow_html=True)

        st.markdown("Tip : Include relevant documents such as market research, technical specifications, user stories, meeting notes, and email trails for a comprehensive BRD.")

        uploaded_files = st.file_uploader("Upload Files", type=["pdf", "docx", "txt"], accept_multiple_files=True)
        
        

        if uploaded_files:
            for uploaded_file in uploaded_files:
                if uploaded_file not in [doc['file'] for doc in st.session_state.support_docs]:
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        doc_description = st.text_input(f"Description for {uploaded_file.name}", key=uploaded_file.name)
                    with col2:
                        if st.button("Add", key=f"add_{uploaded_file.name}"):
                            if doc_description:
                                st.session_state.support_docs.append({
                                    "file": uploaded_file,
                                    "description": doc_description
                                })
                                st.success(f"Added: {uploaded_file.name}")
                            else:
                                st.error("Please provide a description.")

        # Display added documents
        if st.session_state.support_docs:
            st.subheader("Added Documents")
            for idx, doc in enumerate(st.session_state.support_docs):
                st.write(f"{idx + 1}. {doc['file'].name} - {doc['description']}")

        # Process BRD button
        if st.button("Process BRD"):
            if st.session_state.template_file is not None and st.session_state.support_docs:
                # Process each document individually
                with st.spinner("Processing documents..."):
                    for doc in st.session_state.support_docs:
                        processed_doc = process_single_document(
                            file_content=doc["file"].read(),
                            description=doc["description"],
                            document_type=doc["file"].type,
                            session_folder=st.session_state.session_folder
                        )
                        if processed_doc["path"]:
                            st.session_state.processed_docs.append(processed_doc)
                
                # Display processed documents for verification
                display_processed_documents(st.session_state.processed_docs)
                
                if not st.session_state.processed_docs:
                    st.error("No documents were successfully processed. Please check the logs for errors.")
                    return
                
                # Prepare input for next stage
                combined_summary = "\n\n".join([f"Document: {doc['summary'].get('title', 'No title')}\n{doc['summary'].get('description', 'No description')}" for doc in st.session_state.processed_docs])
                template_content = st.session_state.template_file.read()
                
                brd_input = BRDInput(
                    prompt=f"{prompt}\n\nAnalyze the following summaries:\n{combined_summary}",
                    template=template_content,
                    support_documents=[]  # We've already processed the documents
                )


                # Process BRD
                with st.spinner("Processing BRD..."):
                    st.session_state.brd_results = process_brd(brd_input)
                st.session_state.processing_stage = 'completion_toggle'
                st.rerun()
            else:                
                if st.session_state.template_file is None:
                    st.error("Please upload a BRD template file (PDF).")
                if not st.session_state.support_docs:
                    st.error("Please add at least one support document.")

    # elif st.session_state.processing_stage == 'reworded_summary':
    #     st.header("Reworded writeup")
    #     st.write(st.session_state.brd_results["reworded_summary"])

    #     if st.button("Next Process"):
    #         st.session_state.processing_stage = 'completion_toggle'
    #         st.rerun()

    # Completion toggle section
    elif st.session_state.processing_stage == 'completion_toggle':

        st.markdown(f"""
            <style>
                .small-header {{
                    font-size: 20px;
                    font-weight: bold;
                    color: white;
                }}
            </style>
            <div class="small-header">
                Interactive Mode
            </div>
        """, unsafe_allow_html=True)

        # st.header("Provide Additional Information")
        st.write("Please provide additional information for BRD creation?")
        st.session_state.need_completion = st.toggle("Interactive Mode", value=True)

        if st.button("Proceed"):
            if st.session_state.need_completion:
                st.session_state.processing_stage = 'completion'
            else:
                st.session_state.processing_stage = 'generate_final'
            st.rerun()

    # Completion questions section
    elif st.session_state.processing_stage == 'completion':
        # st.header("Provide Additional Information")
        st.markdown(f"""
            <style>
                .small-header {{
                    font-size: 20px;
                    font-weight: bold;
                    color: white;
                }}
            </style>
            <div class="small-header">
                Provide Additional Information
            </div>
        """, unsafe_allow_html=True)

        st.write("Please provide answers to the following questions to complete the BRD:")

        completion_suggestions = st.session_state.brd_results['completion_suggestions']
        if completion_suggestions['status'] == 'need':
            all_questions_answered = True
            for detail in completion_suggestions['details']:
                answer = st.text_input(detail, key=detail)
                if answer:
                    st.session_state.completion_answers[detail] = answer
                else:
                    all_questions_answered = False

            if all_questions_answered:
                if st.button("Proceed to Final BRD Generation"):
                    st.session_state.processing_stage = 'generate_final'
                    st.rerun()
        else:
            st.write("No additional information needed.")
            st.session_state.processing_stage = 'generate_final'
            st.rerun()

    # Generate Final BRD section
    elif st.session_state.processing_stage == 'generate_final':
        st.header("Generating Final BRD")
        with st.spinner("Generating final BRD document..."):
            # Recreate the BRDInput object
            template_content = st.session_state.template_file.read()
            support_documents = [
                SupportDocument(
                    file_content=doc["file"].read(),
                    description=doc["description"],
                    document_type=doc["file"].type
                ) for doc in st.session_state.support_docs
            ]
            brd_input = BRDInput(
                prompt=st.session_state.brd_results["reworded_summary"],
                template=template_content,
                support_documents=support_documents
            )
            reworded_summary = st.session_state.brd_results["reworded_summary"]
            
            # Modify the completion_suggestions based on the toggle
            completion_suggestions = st.session_state.brd_results['completion_suggestions']
            if not st.session_state.need_completion:
                completion_suggestions['status'] = 'not_need'
                completion_suggestions['details'] = []
            
            st.session_state.final_brd = generate_final_brd(brd_input, st.session_state.completion_answers, reworded_summary)
        st.session_state.processing_stage = 'results'
        st.rerun()

    # Results section
    elif st.session_state.processing_stage == 'results':
        # st.header("Generated BRD")

        # st.markdown(f"""
        #     <style>
        #         .small-header {{
        #             font-size: 20px;
        #             font-weight: bold;
        #             color: white;
        #         }}
        #     </style>
        #     <div class="small-header">
        #         Generated BRD
        #     </div>
        # """, unsafe_allow_html=True)

        # st.subheader("Reworded writeup")
        # st.write(st.session_state.brd_results["reworded_summary"])

        # if st.session_state.need_completion:
            # st.subheader("Completion Suggestions and Answers")
            # for question, answer in st.session_state.completion_answers.items():
                # st.write(f"**Q:** {question}")
                # st.write(f"**A:** {answer}")

        # st.subheader("BRD Document")
        # st.text_area("Generated BRD", st.session_state.final_brd["brd_document"], height=600)
        # import streamlit as st


        def markdown_to_html(markdown_text):
            html = markdown.markdown(markdown_text)
            return html

            # Your existing code
        st.subheader("BRD Document")

            # Format the BRD content
        formatted_brd = markdown_to_html(st.session_state.final_brd["brd_document"])
        # word_text = html_to_word(formatted_brd)

            # Display the formatted BRD in a scrollable container
        st.markdown(
                f"""
                <div style="height: 600px; overflow-y: scroll; padding: 10px; border: 1px solid #ccc; border-radius: 5px;">
                    {formatted_brd}
                </div>
                """,
                unsafe_allow_html=True
            )

        # If you want to provide the raw text for copying, you can add this:
        with st.expander("View raw BRD text"):
            st.text_area("Generated BRD", st.session_state.final_brd["brd_document"], height=600)

        #     st.markdown("""
        # <style>
        # .stMarkdown {
        #     font-family: Arial, sans-serif;
        #     padding: 20px;
        #     background-color: #f0f0f0;
        #     border-radius: 5px;
        # }
        # </style>
        # """, unsafe_allow_html=True)
        # Generate Word document
        # output_filename = "brd_document.docx"
        # create_brd_word_document(
        #     st.session_state.final_brd,
        #     output_filename
        # )

        # Option to download BRD as text
        st.download_button(
            label="Download BRD Document (Text)",
            data=formatted_brd,
            file_name="brd_document.txt",
            mime="text/plain"
        )

        # Generate Word document
        output_filename = "Business_Requirements_Document.docx"
        html_to_word(
            formatted_brd,
            output_filename
        )

        # Option to download BRD as Word document
        with open(output_filename, "rb") as file:
            st.download_button(
                label="Download BRD Document (Word)",
                data=file,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    # Instructions
    st.sidebar.header("Instructions")
    st.sidebar.markdown("""
    1. Enter the BRD prompt in the text area.
    2. Upload a BRD template file (PDF).
    3. Upload one or more support documents.
    4. For each support document:
       - Provide a description
       - Click "Add" to include it in the process
    5. Click "Process BRD" when ready.
    6. Choose whether you need to Intractive mode.
    7. If yes, answer the questions provided.
    8. Review the generated BRD and feedback.
    9. Download the final BRD document if desired.
    """)

if __name__ == "__main__":
    main()